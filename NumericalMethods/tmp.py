from docx import Document
from docx.shared import Inches
from docx2pdf import convert

# Создаем документ
doc = Document()
doc.add_heading("Тестирование веб-приложения «Список задач»", level=1)

# Чек-лист
doc.add_heading("1. Чек-лист для ручного тестирования", level=2)
checklist = [
    "Регистрация нового пользователя (позитивный кейс)",
    "Регистрация с уже существующим email (негативный кейс)",
    "Вход с корректными учетными данными",
    "Вход с некорректным паролем (негативный кейс)",
    "Создание новой задачи с заполненными полями",
    "Создание задачи с пустым названием (негативный кейс)",
    "Редактирование задачи и сохранение изменений",
    "Удаление задачи и проверка исчезновения из списка",
    "Фильтрация задач по статусу: все / активные / выполненные",
    "Завершение задачи и проверка фильтрации",
]
for item in checklist:
    doc.add_paragraph(f"- {item}")

doc.add_paragraph()
p = doc.add_paragraph("МЕМО-ПАУЗА: «Не баг, а фича» — сказал кот и ушёл смотреть логи.")
p.italic = True

# Вставляем мем с котиком
meme_url = "https://hips.hearstapps.com/hmg-prod/images/womanyellingcat-1573233850.jpg"  # Замените на реальный URL
doc.add_paragraph("Мем с котиком:")
doc.add_paragraph(f"![Мем с котиком]({meme_url})")

# Баг-репорты
doc.add_heading("2. Баг-репорты (вымышленные)", level=2)
bugs = [
    "При создании задачи с пустым названием возникает 500 ошибка вместо валидации.",
    "После удаления задача визуально остается в списке до перезагрузки страницы.",
    "При неверном пароле alert не содержит текста ошибки.",
    "Фильтр «выполненные» отображает все задачи.",
    "После редактирования данные не сохраняются, хотя приходит сообщение об успехе.",
]
for bug in bugs:
    doc.add_paragraph(f"- {bug}")

doc.add_paragraph()
p = doc.add_paragraph("КОТИК ГОВОРИТ: «Ты не тестер, если ни разу не починил баг… багом же!»")
p.italic = True

# Тестирование API
doc.add_heading("3. Тестирование API", level=2)
api_text = """
/login:
- POST: валидные данные → 200 OK + token
- POST: неверный пароль → 401 Unauthorized

/tasks:
- GET: список задач → 200 OK + JSON
- POST: новая задача → 201 Created
- POST: пустое поле → 400 Bad Request

/tasks/{id}:
- PUT: редактирование задачи → 200 OK
- DELETE: удаление → 204 No Content
- GET: несуществующий ID → 404 Not Found
"""
doc.add_paragraph(api_text.strip())

# Оценка рисков и приоритетов
doc.add_heading("4. Оценка рисков и приоритетов", level=2)
risks = """
Приоритеты:
1. Регистрация и вход (без них недоступны функции).
2. Создание / просмотр задач — основной функционал.
3. Фильтрация — влияет на восприятие задач.

Риски:
- Потеря данных при редактировании.
- Ошибки API.
- Проблемы с фильтрацией — может вводить пользователя в заблуждение.
"""
doc.add_paragraph(risks.strip())

# Сохраняем Word
word_file = "test_tasklist_junior_with_meme.docx"
pdf_file = "test_tasklist_junior_with_meme.pdf"
doc.save(word_file)

# Конвертируем в PDF
convert(word_file, pdf_file)

print(f"Документ с мемом создан: {word_file} и конвертирован в PDF: {pdf_file}")
