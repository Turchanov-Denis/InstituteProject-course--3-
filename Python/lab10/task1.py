from docx import Document

doc = Document()

doc.add_heading('Описание памяти микроконтроллеров ATmega', 0)

doc.add_paragraph(
    'В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:')

doc.add_paragraph(
    '• Флеш-память: используется для хранения скетчей.\n'
    '• ОЗУ (SRAM — static random access memory, статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.\n'
    '• EEPROM (энергонезависимая память): используется для хранения постоянной информации.'
)

table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ATmega168'
hdr_cells[1].text = 'ATmega328'
hdr_cells[2].text = 'ATmega1280'
hdr_cells[3].text = 'ATmega2560'

table.add_row().cells[0].text = '16Кбайт (1 кБ flash-памяти занят загрузчиком)'
table.rows[1].cells[1].text = '32Кбайт'
table.rows[1].cells[2].text = '128Кбайт'
table.rows[1].cells[3].text = '256Кбайт'

table.add_row().cells[0].text = '1Кбайт'
table.rows[2].cells[1].text = '2Кбайт'
table.rows[2].cells[2].text = '8Кбайт'
table.rows[2].cells[3].text = '8Кбайт'

table.add_row().cells[0].text = '512байт'
table.rows[3].cells[1].text = '1024байта'
table.rows[3].cells[2].text = '4Кбайт'
table.rows[3].cells[3].text = '4Кбайт'

doc.add_paragraph(
    'Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100000 операций записи/стирания и 100 лет хранения данных при температуре 25°С.'
)

doc.save('test.docx')
