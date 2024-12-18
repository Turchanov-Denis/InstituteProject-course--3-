from docx import Document
from docx.shared import Inches

doc = Document('test.docx')

doc.add_paragraph('Изображение микроконтроллера ATmega328:')
doc.add_picture('atmega328.png', width=Inches(2.0))

doc.add_paragraph('Изображение микроконтроллера ATmega328 с характеристиками памяти.')

doc.save('test.docx')
